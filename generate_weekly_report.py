#!/usr/bin/env python3
"""
Keep Hammering — Weekly Sales Report Generator
Reads this week's 7 daily notes from Apple Notes (STM > Daily Notes),
parses out customer activity from the Notes section, and writes a
formatted Word document to OneDrive/Work/STM/Weekly calls/

Notes section entry formats supported:
  Format A (same line):  "Republic Parker: Stopped by, saw Destiny..."
  Format B (split line): "Republic Parker:"
                         "Stopped by, saw Destiny..."

Run manually:  python3 ~/Documents/KeepHammering/generate_weekly_report.py
Auto-run:      Shortcuts automation — every Sunday at 8:00 AM
"""

import subprocess
import re
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ────────────────────────────────────────────────────────────────────

OUTPUT_DIR = os.path.expanduser(
    "~/Library/CloudStorage/OneDrive-Personal/Work/STM/Weekly calls"
)

MONTH_NAMES = [
    "January","February","March","April","May","June",
    "July","August","September","October","November","December"
]

# Brand colors
COLOR_GOLD    = RGBColor(0xD4, 0xAF, 0x37)   # Keep Hammering gold
COLOR_BLACK   = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_RED     = RGBColor(0xCC, 0x00, 0x00)
COLOR_LGRAY   = RGBColor(0xF2, 0xF2, 0xF2)   # row shading
COLOR_DGRAY   = RGBColor(0x40, 0x40, 0x40)
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# ── Date Helpers ──────────────────────────────────────────────────────────────

def get_monday():
    today = datetime.today()
    return today - timedelta(days=today.weekday())

def format_date_long(d):
    return f"{d.strftime('%A')}, {d.day} {MONTH_NAMES[d.month - 1]} {d.year}"

def format_date_short(d):
    return f"{d.month:02d}-{d.day:02d}"

# ── Apple Notes Reader ────────────────────────────────────────────────────────

def get_note_plaintext(note_title):
    safe = note_title.replace('"', '\\"')
    script = f'''
tell application "Notes"
    set targetFolder to missing value
    tell default account
        set stmFolder to missing value
        repeat with f in folders
            if name of f is "STM" then
                set stmFolder to f
                exit repeat
            end if
        end repeat
        if stmFolder is not missing value then
            repeat with sf in folders of stmFolder
                if name of sf is "Daily Notes" then
                    set targetFolder to sf
                    exit repeat
                end if
            end repeat
        end if
        if targetFolder is not missing value then
            repeat with n in notes of targetFolder
                if name of n is "{safe}" then
                    return plaintext of n
                end if
            end repeat
        end if
    end tell
end tell
return ""
'''
    result = subprocess.run(["osascript", "-e", script],
                            capture_output=True, text=True, timeout=30)
    return result.stdout.strip()

# ── Note Parser ───────────────────────────────────────────────────────────────

def parse_note(plaintext):
    """
    Two customer entry formats are supported:

    Format B (header line, details below):
        APS:
        4 - 255/60R18 F011894 PO# 2394
        4 - P225/60R17 PO# 2384 HNK1021497

    Format A (everything on one line):
        WCJ: stopped by, tried to catch Matt...

    Detection rules (applied in order):
        1. Line ends with ":"                               → Format B header
        2. Has ":" with prefix ≤ 30 chars, suffix ≥ 20 chars → Format A header
        3. Everything else                                  → sub-item or general note

    Returns:
        customer_entries : list of (customer_name, [note_line, ...]) tuples
        general_notes    : list of strings with no customer context
    """
    lines    = plaintext.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    in_notes = False
    raw_lines = []

    for raw in lines:
        line = raw.strip()
        if line in ("Notes:", "Notes\uff1a"):
            in_notes = True
            continue
        if line in ("Scratch pad:", "Scratch pad\uff1a"):
            break
        if in_notes and line:
            raw_lines.append(line)

    customer_entries = []
    general_notes    = []
    current_customer = None
    current_subitems = []

    def flush():
        if current_customer is not None:
            customer_entries.append(
                (current_customer, [s for s in current_subitems if s])
            )

    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            continue

        # ── Rule 1: Format B header — line ends with ":" ─────────────────────
        if stripped.endswith(":"):
            flush()
            current_customer = stripped[:-1].strip()
            current_subitems = []

        # ── Rule 2: Format A header — "Customer: long activity note" ─────────
        elif ":" in stripped:
            parts  = stripped.split(":", 1)
            prefix = parts[0].strip()
            suffix = parts[1].strip()
            if len(prefix) <= 30 and len(suffix) >= 20:
                flush()
                current_customer = prefix
                current_subitems = [suffix]
            else:
                # Short suffix (phone, address field, etc.) → sub-item
                if current_customer is not None:
                    current_subitems.append(stripped)
                else:
                    general_notes.append(stripped)

        # ── Rule 3: plain line → sub-item of current customer ─────────────────
        else:
            if current_customer is not None:
                current_subitems.append(stripped)
            else:
                general_notes.append(stripped)

    flush()  # capture the last customer
    return customer_entries, general_notes

# ── Word Doc Helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Fill a table cell with a solid background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Add borders to a table cell. kwargs: top, bottom, left, right = hex color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in kwargs.items():
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color
    return run

# ── Report Builder ────────────────────────────────────────────────────────────

def build_word_report(week_data, monday):
    sunday  = monday + timedelta(days=6)
    mon_str = f"{MONTH_NAMES[monday.month-1]} {monday.day}"
    sun_str = f"{MONTH_NAMES[sunday.month-1]} {sunday.day}, {sunday.year}"

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Document Header ───────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title_para, "WEEKLY SALES REPORT",
            bold=True, size=20, color=COLOR_BLACK)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub_para, f"Week of {mon_str} – {sun_str}",
            size=12, color=COLOR_DGRAY)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta_para, "Name: ", bold=True, size=11)
    add_run(meta_para, "Vaughan", italic=True, size=11)

    motto_para = doc.add_paragraph()
    motto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(motto_para, "Aut Viam Inveniam Aut Faciam",
            italic=True, bold=True, size=10, color=COLOR_RED)

    doc.add_paragraph()  # spacer

    # ── Day Sections ──────────────────────────────────────────────────────────
    for day_date, customer_entries, general_notes in week_data:
        day_name = day_date.strftime("%A").upper()

        # Day header table (full-width dark band)
        hdr_tbl = doc.add_table(rows=1, cols=1)
        hdr_tbl.style = "Table Grid"
        hdr_cell = hdr_tbl.cell(0, 0)
        set_cell_bg(hdr_cell, "1A1A1A")
        hdr_para = hdr_cell.paragraphs[0]
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hdr_para.paragraph_format.space_before = Pt(4)
        hdr_para.paragraph_format.space_after  = Pt(4)
        add_run(hdr_para, f"  {day_name}",
                bold=True, size=13, color=COLOR_GOLD)

        doc.add_paragraph()  # small gap

        if customer_entries or general_notes:
            # Content table: Customer | Note
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Table Grid"
            tbl.columns[0].width = Inches(1.8)
            tbl.columns[1].width = Inches(4.7)

            # Column header row
            col_row = tbl.add_row()
            col_row.cells[0].width = Inches(1.8)
            col_row.cells[1].width = Inches(4.7)
            set_cell_bg(col_row.cells[0], "D4AF37")
            set_cell_bg(col_row.cells[1], "D4AF37")

            ch0 = col_row.cells[0].paragraphs[0]
            ch1 = col_row.cells[1].paragraphs[0]
            add_run(ch0, "Customer", bold=True, size=10, color=COLOR_BLACK)
            add_run(ch1, "Activity / Notes", bold=True, size=10, color=COLOR_BLACK)

            # Customer rows
            for idx, (customer, note_lines) in enumerate(customer_entries):
                row = tbl.add_row()
                row.cells[0].width = Inches(1.8)
                row.cells[1].width = Inches(4.7)

                bg = "F9F9F9" if idx % 2 == 0 else "FFFFFF"
                set_cell_bg(row.cells[0], bg)
                set_cell_bg(row.cells[1], bg)

                c_para = row.cells[0].paragraphs[0]
                c_para.paragraph_format.space_before = Pt(3)
                c_para.paragraph_format.space_after  = Pt(3)
                add_run(c_para, customer, bold=True, size=10)

                # Each sub-item on its own line inside the Activity cell
                note_cell = row.cells[1]
                if note_lines:
                    for j, note_line in enumerate(note_lines):
                        if j == 0:
                            n_para = note_cell.paragraphs[0]
                        else:
                            n_para = note_cell.add_paragraph()
                        n_para.paragraph_format.space_before = Pt(1)
                        n_para.paragraph_format.space_after  = Pt(1)
                        add_run(n_para, note_line, size=10)
                else:
                    note_cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                    note_cell.paragraphs[0].paragraph_format.space_after  = Pt(3)

            # General notes (no customer prefix)
            for note in general_notes:
                row = tbl.add_row()
                set_cell_bg(row.cells[0], "FFF8E7")
                set_cell_bg(row.cells[1], "FFF8E7")
                add_run(row.cells[0].paragraphs[0],
                        "General", italic=True, size=9, color=COLOR_DGRAY)
                add_run(row.cells[1].paragraphs[0], note, size=10)

        else:
            no_para = doc.add_paragraph()
            add_run(no_para, "  No activity recorded.",
                    italic=True, size=10, color=RGBColor(0xAA, 0xAA, 0xAA))

        doc.add_paragraph()  # spacer between days

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    now = datetime.now()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_para,
            f"Generated: {MONTH_NAMES[now.month-1]} {now.day}, {now.year}  |  Keep Hammering.",
            italic=True, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

    return doc

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    import sys
    offset = -7 if "--last-week" in sys.argv else 0
    monday    = get_monday() + timedelta(days=offset)
    week_data = []

    print("\n  Keep Hammering — Weekly Report Generator")
    print(f"  Week of {format_date_long(monday)}\n")

    for i in range(7):
        day   = monday + timedelta(days=i)
        title = format_date_long(day)
        print(f"  [{i+1}/7] {title} ...", end=" ", flush=True)

        plaintext = get_note_plaintext(title)
        if plaintext:
            customer_entries, general_notes = parse_note(plaintext)
            print(f"{len(customer_entries)} customer(s)")
        else:
            customer_entries, general_notes = [], []
            print("note not found")

        week_data.append((day, customer_entries, general_notes))

    doc = build_word_report(week_data, monday)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = (
        f"Weekly_Report_{format_date_short(monday)}_to_"
        f"{format_date_short(monday + timedelta(days=6))}_{monday.year}.docx"
    )
    out_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(out_path)

    print(f"\n  Saved to: {out_path}\n")

if __name__ == "__main__":
    main()
